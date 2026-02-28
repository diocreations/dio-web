"""
Resume Builder Routes
Create resumes from scratch with AI assistance
"""
from fastapi import APIRouter, HTTPException, Depends
from database import db, EMERGENT_LLM_KEY, logger
from helpers import get_current_user
from datetime import datetime, timezone
from emergentintegrations.llm.chat import LlmChat, UserMessage
import uuid
import asyncio
import io
import os

router = APIRouter(prefix="/api/builder")


# ==================== RESUME DRAFTS ====================

@router.post("/draft")
async def save_draft(data: dict, user: dict = Depends(get_current_user)):
    """Save or update a resume draft"""
    draft_id = data.get("draft_id") or f"draft_{uuid.uuid4().hex[:12]}"
    user_id = user.get("user_id") or user.get("email")
    
    draft_doc = {
        "draft_id": draft_id,
        "user_id": user_id,
        "title": data.get("title", "Untitled Resume"),
        "template_id": data.get("template_id", "classic"),
        "sections": data.get("sections", {}),
        "personal_info": data.get("personal_info", {}),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    
    # Check if draft exists
    existing = await db.resume_drafts.find_one({"draft_id": draft_id, "user_id": user_id})
    if existing:
        await db.resume_drafts.update_one(
            {"draft_id": draft_id},
            {"$set": draft_doc}
        )
    else:
        draft_doc["created_at"] = datetime.now(timezone.utc).isoformat()
        await db.resume_drafts.insert_one(draft_doc)
    
    return {"draft_id": draft_id, "message": "Draft saved successfully"}


@router.get("/drafts")
async def get_user_drafts(user: dict = Depends(get_current_user)):
    """Get all drafts for current user"""
    user_id = user.get("user_id") or user.get("email")
    drafts = await db.resume_drafts.find(
        {"user_id": user_id},
        {"_id": 0, "sections": 0}  # Exclude heavy data for listing
    ).sort("updated_at", -1).to_list(50)
    return drafts


@router.get("/draft/{draft_id}")
async def get_draft(draft_id: str, user: dict = Depends(get_current_user)):
    """Get a specific draft"""
    user_id = user.get("user_id") or user.get("email")
    draft = await db.resume_drafts.find_one(
        {"draft_id": draft_id, "user_id": user_id},
        {"_id": 0}
    )
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft


@router.delete("/draft/{draft_id}")
async def delete_draft(draft_id: str, user: dict = Depends(get_current_user)):
    """Delete a draft"""
    user_id = user.get("user_id") or user.get("email")
    result = await db.resume_drafts.delete_one({"draft_id": draft_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Draft not found")
    return {"message": "Draft deleted"}


# ==================== AI GENERATION ====================

@router.post("/generate/summary")
async def generate_summary(data: dict):
    """Generate professional summary from basic info"""
    job_title = data.get("job_title", "")
    years_exp = data.get("years_experience", "")
    skills = data.get("skills", [])
    achievements = data.get("achievements", "")
    
    if not job_title:
        raise HTTPException(status_code=400, detail="Job title required")
    
    skills_str = ", ".join(skills) if isinstance(skills, list) else skills
    
    prompt = f"""Generate a compelling professional summary for a resume. Keep it 3-4 sentences, impactful and ATS-friendly.

Job Title: {job_title}
Years of Experience: {years_exp or 'Not specified'}
Key Skills: {skills_str or 'Not specified'}
Key Achievements: {achievements or 'Not specified'}

Write in first person implied (no "I"), professional tone. Focus on value proposition and measurable impact where possible."""

    try:
        session_id = f"builder_{uuid.uuid4().hex[:8]}"
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message="You are a professional resume writer.").with_model("gemini", "gemini-2.0-flash")
        response = await chat.send_message(UserMessage(text=prompt))
        return {"summary": response.content.strip()}
    except Exception as e:
        logger.error(f"AI summary generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate summary")


@router.post("/generate/experience")
async def generate_experience_bullets(data: dict):
    """Generate bullet points for work experience"""
    job_title = data.get("job_title", "")
    company = data.get("company", "")
    responsibilities = data.get("responsibilities", "")
    achievements = data.get("achievements", "")
    
    if not job_title:
        raise HTTPException(status_code=400, detail="Job title required")
    
    prompt = f"""Generate 4-6 impactful bullet points for this work experience. Use action verbs, quantify achievements where possible, and make them ATS-friendly.

Job Title: {job_title}
Company: {company or 'Not specified'}
Key Responsibilities: {responsibilities or 'General duties'}
Achievements: {achievements or 'Not specified'}

Format: Return ONLY the bullet points, one per line, starting with action verbs (Led, Developed, Managed, Increased, etc.). No bullet characters, just the text."""

    try:
        session_id = f"builder_{uuid.uuid4().hex[:8]}"
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message="You are a professional resume writer.").with_model("gemini", "gemini-2.0-flash")
        response = await chat.send_message(UserMessage(text=prompt))
        bullets = [b.strip() for b in response.content.strip().split("\n") if b.strip()]
        return {"bullets": bullets}
    except Exception as e:
        logger.error(f"AI experience generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate experience")


@router.post("/generate/skills")
async def generate_skills(data: dict):
    """Generate relevant skills based on job title"""
    job_title = data.get("job_title", "")
    industry = data.get("industry", "")
    current_skills = data.get("current_skills", [])
    
    if not job_title:
        raise HTTPException(status_code=400, detail="Job title required")
    
    current_str = ", ".join(current_skills) if current_skills else "None provided"
    
    prompt = f"""Suggest relevant skills for this role. Include both technical and soft skills that would be valuable.

Job Title: {job_title}
Industry: {industry or 'General'}
Current Skills: {current_str}

Return as JSON with two arrays:
{{"technical": ["skill1", "skill2", ...], "soft": ["skill1", "skill2", ...]}}

Suggest 8-10 technical skills and 5-6 soft skills. Only return the JSON, nothing else."""

    try:
        session_id = f"builder_{uuid.uuid4().hex[:8]}"
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message="You are a professional resume writer.").with_model("gemini", "gemini-2.0-flash")
        response = await chat.send_message(UserMessage(text=prompt))
        import json
        # Clean response and parse JSON
        content = response.content.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        skills = json.loads(content)
        return skills
    except Exception as e:
        logger.error(f"AI skills generation failed: {e}")
        return {"technical": [], "soft": []}


@router.post("/generate/full")
async def generate_full_resume(data: dict):
    """Generate a complete resume from minimal input"""
    job_title = data.get("job_title", "")
    years_exp = data.get("years_experience", "")
    industry = data.get("industry", "")
    skills = data.get("skills", [])
    education = data.get("education", "")
    
    if not job_title:
        raise HTTPException(status_code=400, detail="Job title required")
    
    skills_str = ", ".join(skills) if isinstance(skills, list) else skills
    edu_default = "Bachelors degree"
    
    prompt = f"""Generate a complete professional resume content for:

Job Title: {job_title}
Years of Experience: {years_exp or '3-5 years'}
Industry: {industry or 'Technology'}
Key Skills: {skills_str or 'Relevant to the role'}
Education: {education or edu_default}

Return as JSON with this structure:
{{
  "summary": "Professional summary text (3-4 sentences)",
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, Country",
      "start_date": "Month Year",
      "end_date": "Present or Month Year",
      "bullets": ["Achievement 1", "Achievement 2", "Achievement 3", "Achievement 4"]
    }}
  ],
  "skills": {{
    "technical": ["skill1", "skill2"],
    "soft": ["skill1", "skill2"]
  }},
  "education": [
    {{
      "degree": "Degree Name",
      "school": "University Name",
      "location": "City, Country",
      "year": "Year"
    }}
  ]
}}

Generate 2-3 relevant work experiences. Make it realistic and ATS-friendly. Only return valid JSON."""

    try:
        session_id = f"builder_{uuid.uuid4().hex[:8]}"
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message="You are a professional resume writer.").with_model("gemini", "gemini-2.0-flash")
        response = await asyncio.to_thread(
            chat.send_message,
            UserMessage(text=prompt)
        )
        import json
        content = response.content.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        resume_data = json.loads(content)
        return resume_data
    except Exception as e:
        logger.error(f"AI full resume generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate resume")


# ==================== EXPORT ====================

@router.post("/export/docx")
async def export_docx(data: dict):
    """Export resume as DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import base64
        
        personal = data.get("personal_info", {})
        sections = data.get("sections", {})
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Name
        name = personal.get("name", "Your Name")
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(22)
        name_run.font.color.rgb = RGBColor(26, 26, 46)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if personal.get("email"): contact_parts.append(personal["email"])
        if personal.get("phone"): contact_parts.append(personal["phone"])
        if personal.get("location"): contact_parts.append(personal["location"])
        if personal.get("linkedin"): contact_parts.append(personal["linkedin"])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.add_run(" | ".join(contact_parts)).font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Summary
        if sections.get("summary"):
            h = doc.add_paragraph()
            h_run = h.add_run("PROFESSIONAL SUMMARY")
            h_run.bold = True
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor(26, 26, 46)
            
            p = doc.add_paragraph()
            p.add_run(sections["summary"]).font.size = Pt(10)
        
        # Experience
        if sections.get("experience"):
            h = doc.add_paragraph()
            h_run = h.add_run("WORK EXPERIENCE")
            h_run.bold = True
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor(26, 26, 46)
            
            for exp in sections["experience"]:
                # Title and dates
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
                title_run.bold = True
                title_run.font.size = Pt(10)
                
                date_para = doc.add_paragraph()
                date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', '')} | {exp.get('location', '')}"
                date_run = date_para.add_run(date_text)
                date_run.italic = True
                date_run.font.size = Pt(9)
                
                # Bullets
                for bullet in exp.get("bullets", []):
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.add_run(bullet).font.size = Pt(10)
        
        # Education
        if sections.get("education"):
            h = doc.add_paragraph()
            h_run = h.add_run("EDUCATION")
            h_run.bold = True
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor(26, 26, 46)
            
            for edu in sections["education"]:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu.get('degree', '')} - {edu.get('school', '')}")
                edu_run.bold = True
                edu_run.font.size = Pt(10)
                
                if edu.get("year") or edu.get("location"):
                    detail_para = doc.add_paragraph()
                    detail_para.add_run(f"{edu.get('year', '')} | {edu.get('location', '')}").font.size = Pt(9)
        
        # Skills
        if sections.get("skills"):
            h = doc.add_paragraph()
            h_run = h.add_run("SKILLS")
            h_run.bold = True
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor(26, 26, 46)
            
            skills = sections["skills"]
            if isinstance(skills, dict):
                if skills.get("technical"):
                    tech_para = doc.add_paragraph()
                    tech_para.add_run("Technical: ").bold = True
                    tech_para.add_run(", ".join(skills["technical"])).font.size = Pt(10)
                if skills.get("soft"):
                    soft_para = doc.add_paragraph()
                    soft_para.add_run("Soft Skills: ").bold = True
                    soft_para.add_run(", ".join(skills["soft"])).font.size = Pt(10)
            elif isinstance(skills, list):
                skills_para = doc.add_paragraph()
                skills_para.add_run(", ".join(skills)).font.size = Pt(10)
        
        # Certifications
        if sections.get("certifications"):
            h = doc.add_paragraph()
            h_run = h.add_run("CERTIFICATIONS")
            h_run.bold = True
            h_run.font.size = Pt(11)
            h_run.font.color.rgb = RGBColor(26, 26, 46)
            
            for cert in sections["certifications"]:
                cert_para = doc.add_paragraph(style='List Bullet')
                cert_text = cert.get("name", cert) if isinstance(cert, dict) else cert
                cert_para.add_run(cert_text).font.size = Pt(10)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return as base64
        docx_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        filename = f"{personal.get('name', 'resume').replace(' ', '_').lower()}_resume.docx"
        
        return {
            "docx_base64": docx_base64,
            "filename": filename
        }
        
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX export not available. Please install python-docx.")
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to export DOCX: {str(e)}")


# ==================== PRICING ====================

@router.get("/pricing")
async def get_builder_pricing():
    """Get resume builder pricing settings"""
    pricing = await db.builder_pricing.find_one({"pricing_id": "resume_builder"}, {"_id": 0})
    if not pricing:
        # Default pricing
        pricing = {
            "pricing_id": "resume_builder",
            "enabled": False,  # Free by default
            "price": 4.99,
            "currency": "EUR",
            "product_name": "Resume Builder Pro",
            "description": "Create professional resumes with AI assistance"
        }
    return pricing


@router.put("/pricing")
async def update_builder_pricing(data: dict, user: dict = Depends(get_current_user)):
    """Update resume builder pricing (admin only)"""
    pricing_doc = {
        "pricing_id": "resume_builder",
        "enabled": data.get("enabled", False),
        "price": data.get("price", 4.99),
        "currency": data.get("currency", "EUR"),
        "product_name": data.get("product_name", "Resume Builder Pro"),
        "description": data.get("description", ""),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.builder_pricing.update_one(
        {"pricing_id": "resume_builder"},
        {"$set": pricing_doc},
        upsert=True
    )
    return {"message": "Pricing updated"}
